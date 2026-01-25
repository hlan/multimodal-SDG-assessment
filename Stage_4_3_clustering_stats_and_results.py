# -*- coding: utf-8 -*-

import pandas as pd
import numpy as np
from pathlib import Path

OUTDIR = Path("cluster_outputs")
ASSIGN_CSV = OUTDIR / "cluster_assignments_k7.csv"
SDG_FEAT_CSV = OUTDIR / "features_sdg.csv"
ANOVA_CSV = OUTDIR / "anova_sdg_cluster.csv"
MANOVA_TXT = OUTDIR / "manova_results.txt"

#optional
USE_QUANTILE_BINNING = False   

# optional, not in use
Z_HI2 = 0.80
Z_HI1 = 0.30
Z_LO1 = -0.30
Z_LO2 = -0.80

# optional, not in use
Q_LO2 = 0.20
Q_LO1 = 0.40
Q_HI1 = 0.60
Q_HI2 = 0.80

SDG_NAMES = {
    1: "No Poverty", 2: "Zero Hunger", 3: "Good Health and Well-Being",
    4: "Quality Education", 5: "Gender Equality", 6: "Clean Water and Sanitation",
    7: "Affordable and Clean Energy", 8: "Decent Work and Economic Growth",
    9: "Industry, Innovation and Infrastructure", 10: "Reduced Inequalities",
    11: "Sustainable Cities and Communities", 12: "Responsible Consumption and Production",
    13: "Climate Action", 14: "Life Below Water", 15: "Life on Land",
    16: "Peace, Justice and Strong Institutions", 17: "Partnerships for the Goals"
}

assign = pd.read_csv(ASSIGN_CSV)
sdg_feat = pd.read_csv(SDG_FEAT_CSV)
sdg_mean_cols = [f"sdg_{i}__mean_all" for i in range(1,18)]
df = assign.merge(sdg_feat[["countrycode","countryname"]+sdg_mean_cols],
                  on=["countrycode","countryname"], how="left")

# =====  ANOVA/MANOVA =====
if not ANOVA_CSV.exists() or not MANOVA_TXT.exists():
    import statsmodels.api as sm
    from statsmodels.formula.api import ols
    from statsmodels.multivariate.manova import MANOVA
    formula = " + ".join(sdg_mean_cols) + " ~ C(cluster)"
    maov = MANOVA.from_formula(formula, data=df)
    with open(MANOVA_TXT, "w", encoding="utf-8") as f:
        f.write(str(maov.mv_test()))
    rows = []
    for c in sdg_mean_cols:
        mdl = ols(f"{c} ~ C(cluster)", data=df).fit()
        aov = sm.stats.anova_lm(mdl, typ=2)
        ss_between = aov.loc["C(cluster)","sum_sq"]
        ss_total   = ss_between + aov.loc["Residual","sum_sq"]
        eta2 = ss_between / ss_total if ss_total > 0 else np.nan
        pval = aov.loc["C(cluster)","PR(>F)"]
        rows.append({"sdg":c.replace("__mean_all","").upper(), "eta2":eta2, "pval":pval})
    anova_tbl = pd.DataFrame(rows).sort_values("eta2", ascending=False)
    anova_tbl.to_csv(ANOVA_CSV, index=False)
else:
    anova_tbl = pd.read_csv(ANOVA_CSV)

# ===== Cluster × SDG Z-score matrix =====
cluster_means = df.groupby("cluster")[sdg_mean_cols].mean()
cluster_means_z = (cluster_means - cluster_means.mean()) / cluster_means.std(ddof=0)
report = cluster_means_z.copy()
report.columns = [c.replace("__mean_all","").upper() for c in report.columns]  # SDG_1..SDG_17

# ===== sig? =====
sig_sdg = {row["sdg"]: (row["pval"] < 0.05) for _, row in anova_tbl.iterrows()}

# ===== heatmap preview =====
try:
    import seaborn as sns
    import matplotlib.pyplot as plt
    plt.figure(figsize=(12,6))
    sns.heatmap(
        report, annot=True, fmt=".2f", cmap="coolwarm",
        xticklabels=[(c+"*" if sig_sdg.get(c,False) else c) for c in report.columns],
        yticklabels=[f"Cluster {i}" for i in report.index]
    )
    plt.title("Cluster × SDG (mean_all, z-score)  *p<0.05")
    plt.tight_layout()
    plt.savefig(OUTDIR/"cluster_sdg_report_heatmap.png", dpi=220)
    plt.close()
    print("output：cluster_sdg_report_heatmap.png")
except Exception as e:
    print("fail, missing seaborn/matplotlib：", e)

# ===== labeling =====
labels = {}
for k in report.index:
    zc = report.loc[k].sort_values(ascending=False)
    highs = [col for col in zc.index if sig_sdg.get(col, False)][:2] or zc.head(2).index.tolist()
    lows_all = [col for col in zc.index if sig_sdg.get(col, False)]
    lows  = lows_all[-1:] if len(lows_all) else zc.tail(1).index.tolist()
    labels[k] = f"High {', '.join(highs)}; Low {', '.join(lows)}"
assign_labeled = assign.copy()
assign_labeled["cluster_label_auto"] = assign_labeled["cluster"].map(labels)

# ===== 7 Level =====
LEVEL_ORDER = ["EL", "L", "ML", "M", "MH", "H", "EH"]  # Extreme Low → Extreme High
LABEL_LONG = {
    "EH": "Extreme High", "H": "High", "MH": "Mid-High",
    "M": "Medium", "ML": "Mid-Low", "L": "Low", "EL": "Extreme Low"
}

def level_from_z_fixed(z: float) -> str:
    # EH: z >= 1.5
    # H : 1.0 <= z < 1.5
    # MH: 0.5 <= z < 1.0
    # M : 0.0 <= z < 0.5
    # ML: -0.75 <= z < 0.0
    # L : -1.5 <= z < -0.75
    # EL: z < -1.5
    if z >= 1.5:
        return "EH"
    elif z >= 1.0:
        return "H"
    elif z >= 0.5:
        return "MH"
    elif z >= 0.0:
        return "M"
    elif z >= -0.75:
        return "ML"
    elif z >= -1.5:
        return "L"
    else:
        return "EL"

if USE_QUANTILE_BINNING:
    # optional
    qs = [1/7, 2/7, 3/7, 4/7, 5/7, 6/7]
    q_map = {q: report.quantile(q, axis=0) for q in qs}

    def level_from_quantile(z, col):
        b1 = q_map[1/7][col]
        b2 = q_map[2/7][col]
        b3 = q_map[3/7][col]
        b4 = q_map[4/7][col]
        b5 = q_map[5/7][col]
        b6 = q_map[6/7][col]
        if z < b1:
            return "EL"
        elif z < b2:
            return "L"
        elif z < b3:
            return "ML"
        elif z < b4:
            return "M"
        elif z < b5:
            return "MH"
        elif z < b6:
            return "H"
        else:
            return "EH"

    level_str_df = report.copy()
    for col in report.columns:
        level_str_df[col] = report[col].apply(lambda z: level_from_quantile(z, col))
else:
    level_str_df = report.applymap(level_from_z_fixed)

detailed_rows = []
for k in report.index:
    for col in report.columns:
        z = report.loc[k, col]
        lvl = level_str_df.loc[k, col]
        sdg_num = int(col.split("_")[1])
        detailed_rows.append({
            "cluster": k, "sdg": col,
            "sdg_name": SDG_NAMES[sdg_num],
            "z_score": z, "level_7band": lvl
        })
detailed_out = pd.DataFrame(detailed_rows)

summary_rows = []
for k in report.index:
    bands = {b: [] for b in LEVEL_ORDER}
    for col in report.columns:
        lvl = level_str_df.loc[k, col]
        name = SDG_NAMES[int(col.split("_")[1])]
        bands[lvl].append(name)
    def j(lst): return ", ".join(lst) if lst else "None"
    row = {"cluster": k,
           "summary_en": f"{LABEL_LONG['EH']}: {j(bands['EH'])}; "
                         f"{LABEL_LONG['H']}: {j(bands['H'])}; "
                         f"{LABEL_LONG['MH']}: {j(bands['MH'])}; "
                         f"{LABEL_LONG['M']}: {j(bands['M'])}; "
                         f"{LABEL_LONG['ML']}: {j(bands['ML'])}; "
                         f"{LABEL_LONG['L']}: {j(bands['L'])}; "
                         f"{LABEL_LONG['EL']}: {j(bands['EL'])}"}
    for col in report.columns:
        row[f"{col}_LEVEL"] = level_str_df.loc[k, col]
    summary_rows.append(row)
summary_out = pd.DataFrame(summary_rows).sort_values("cluster")

from statsmodels.formula.api import ols
import statsmodels.api as sm
from statsmodels.stats.multicomp import pairwise_tukeyhsd

# 95% CI 
try:
    from scipy import stats
    def t_crit(n): 
        df_t = max(n-1, 1)
        return stats.t.ppf(1-0.025, df_t)
except Exception:
    def t_crit(n): 
        return 1.96

desc_frames = []
anova_full_frames = []
tukey_frames = []

for i in range(1, 18):
    col_raw = f"sdg_{i}__mean_all"
    sdg_code = f"SDG_{i}"
    # Descriptives
    g = df.groupby("cluster")[col_raw].agg(['mean','std','count']).reset_index()
    g["sdg"] = sdg_code
    # CI
    g["sem"] = g["std"] / g["count"].replace(0, np.nan)**0.5
    g["tcrit"] = g["count"].apply(t_crit)
    g["ci_low"]  = g["mean"] - g["tcrit"] * g["sem"]
    g["ci_high"] = g["mean"] + g["tcrit"] * g["sem"]
    desc_frames.append(g[["sdg","cluster","mean","std","count","ci_low","ci_high"]])

    # ANOVA Full
    mdl = ols(f"{col_raw} ~ C(cluster)", data=df).fit()
    aov = sm.stats.anova_lm(mdl, typ=2).reset_index().rename(columns={"index":"source"})
    aov["sdg"] = sdg_code
    # MS
    aov["mean_sq"] = aov["sum_sq"] / aov["df"]
    anova_full_frames.append(aov[["sdg","source","sum_sq","df","mean_sq","F","PR(>F)"]])

    # Tukey HSD optional
    try:
        tuk = pairwise_tukeyhsd(endog=df[col_raw], groups=df["cluster"], alpha=0.05)
        tdf = pd.DataFrame(tuk.summary().data[1:], columns=tuk.summary().data[0])
        tdf.insert(0, "sdg", sdg_code)
        tukey_frames.append(tdf)
    except Exception:
        pass

descriptives_all = pd.concat(desc_frames, ignore_index=True) if desc_frames else pd.DataFrame()
anova_full_all   = pd.concat(anova_full_frames, ignore_index=True) if anova_full_frames else pd.DataFrame()
tukey_all        = pd.concat(tukey_frames, ignore_index=True) if tukey_frames else pd.DataFrame()

excel_path = OUTDIR / "cluster_sdg_results.xlsx"
with pd.ExcelWriter(excel_path, engine="openpyxl") as writer:
    report.to_excel(writer, sheet_name="Cluster_SDG_Zscores")
    anova_tbl.to_excel(writer, sheet_name="ANOVA_Summary", index=False)
    if not anova_full_all.empty:
        anova_full_all.to_excel(writer, sheet_name="ANOVA_Full", index=False)
    if not descriptives_all.empty:
        descriptives_all.to_excel(writer, sheet_name="Descriptives", index=False)
    if not tukey_all.empty:
        tukey_all.to_excel(writer, sheet_name="Tukey_HSD", index=False)
    assign_labeled.to_excel(writer, sheet_name="Assignments_Labeled", index=False)
    detailed_out.to_excel(writer, sheet_name="SDG_Levels_Detailed_7band", index=False)
    summary_out.to_excel(writer, sheet_name="SDG_Levels_Summary_7band", index=False)

print(f"save to Excel：{excel_path}")
print("Sheets: Cluster_SDG_Zscores, ANOVA_Summary, ANOVA_Full, Descriptives, Tukey_HSD, "
      "Assignments_Labeled, SDG_Levels_Detailed_7band, SDG_Levels_Summary_7band")

# ===== optional output =====
DESC_DOCX = OUTDIR / "ANOVA_Descriptives.docx"
ANOVA_DOCX = OUTDIR / "ANOVA_Table.docx"
TUKEY_DOCX = OUTDIR / "ANOVA_Multiple_Comparisons_Tukey.docx"

DESC_XLSX = OUTDIR / "ANOVA_Descriptives.xlsx"
ANOVA_XLSX = OUTDIR / "ANOVA_Table.xlsx"
TUKEY_XLSX = OUTDIR / "ANOVA_Multiple_Comparisons_Tukey.xlsx"

def tidy_descriptives(df_desc: pd.DataFrame) -> pd.DataFrame:
    if df_desc.empty: 
        return df_desc
    d = df_desc.copy()
    d = d.rename(columns={
        "cluster":"Cluster", "mean":"Mean", "std":"Std. Deviation",
        "count":"N", "ci_low":"95% CI Lower", "ci_high":"95% CI Upper", "sdg":"SDG"
    })
    # SDG、Cluster
    d = d[["SDG","Cluster","N","Mean","Std. Deviation","95% CI Lower","95% CI Upper"]]
    for c in ["Mean","Std. Deviation","95% CI Lower","95% CI Upper"]:
        d[c] = d[c].astype(float).round(3)
    d["Cluster"] = d["Cluster"].astype(int)
    return d

def tidy_anova_full(df_anova: pd.DataFrame) -> pd.DataFrame:
    if df_anova.empty:
        return df_anova
    a = df_anova.copy()
    # Between Groups / Within Groups
    a["source"] = a["source"].replace({"C(cluster)":"Between Groups", "Residual":"Within Groups"})
    a = a[a["source"].isin(["Between Groups","Within Groups"])].copy()
    total = a.groupby("sdg", as_index=False).agg({"sum_sq":"sum","df":"sum"})
    total["source"] = "Total"
    total["mean_sq"] = np.nan
    total["F"] = np.nan
    total["PR(>F)"] = np.nan
    a = pd.concat([a, total[["sdg","source","sum_sq","df","mean_sq","F","PR(>F)"]]], ignore_index=True)

    a = a.rename(columns={
        "sdg":"SDG","source":"Source","sum_sq":"Sum of Squares",
        "df":"df","mean_sq":"Mean Square","F":"F","PR(>F)":"Sig."
    })
    # ranking
    order_map = {"Between Groups":0, "Within Groups":1, "Total":2}
    a["order"] = a["Source"].map(order_map)
    a = a.sort_values(["SDG","order"]).drop(columns=["order"])
    for c in ["Sum of Squares","Mean Square","F","Sig."]:
        a[c] = a[c].astype(float).round(4)
    a["df"] = a["df"].astype(float).round(0).astype("Int64")
    return a

def tidy_tukey(df_tukey: pd.DataFrame) -> pd.DataFrame:
    if df_tukey.empty:
        return df_tukey
    t = df_tukey.copy()
    rename_map = {
        "group1":"Group I", "group2":"Group J",
        "meandiff":"Mean Difference (I-J)",
        "p-adj":"Sig.", "lower":"Lower Bound", "upper":"Upper Bound",
        "reject":"Significant?"
    }
    t = t.rename(columns=rename_map)
    keep = ["sdg","Group I","Group J","Mean Difference (I-J)","Sig.","Lower Bound","Upper Bound","Significant?"]
    t = t[keep].rename(columns={"sdg":"SDG"})
    t["Mean Difference (I-J)"] = t["Mean Difference (I-J)"].astype(float).round(4)
    for c in ["Sig.","Lower Bound","Upper Bound"]:
        t[c] = t[c].astype(float).round(4)
    t["Significant?"] = t["Significant?"].map({True:"Yes", False:"No"})
    for c in ["Group I","Group J"]:
        t[c] = t[c].astype(float).round(0).astype(int).astype(str)
    t = t.sort_values(["SDG","Group I","Group J"]).reset_index(drop=True)
    return t

desc_tidy  = tidy_descriptives(descriptives_all) if 'descriptives_all' in locals() else pd.DataFrame()
anova_tidy = tidy_anova_full(anova_full_all)     if 'anova_full_all' in locals() else pd.DataFrame()
tukey_tidy = tidy_tukey(tukey_all)               if 'tukey_all' in locals() else pd.DataFrame()

def write_docx_table(filepath: Path, df: pd.DataFrame, title: str, subtitle_col: str = "SDG"):
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    h = doc.add_heading(title, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if df.empty:
        p = doc.add_paragraph("No data available.")
        doc.save(str(filepath))
        return

    cols = df.columns.tolist()
    group_key = subtitle_col if subtitle_col in df.columns else None
    for sdg, g in (df.groupby(group_key) if group_key else [(None, df)]):
        if group_key:
            doc.add_heading(str(sdg), level=2)
        table = doc.add_table(rows=1, cols=len(cols))
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        for j, c in enumerate(cols):
            hdr_cells[j].text = c
        for _, row in g.iterrows():
            cells = table.add_row().cells
            for j, c in enumerate(cols):
                val = row[c]
                cells[j].text = "" if pd.isna(val) else str(val)
        doc.add_paragraph("")

    sections = doc.sections
    for section in sections:
        section.left_margin  = Inches(0.6)
        section.right_margin = Inches(0.6)
        section.top_margin   = Inches(0.6)
        section.bottom_margin= Inches(0.6)
    doc.save(str(filepath))

def export_three_docs():
    try:
        import docx  
        if not desc_tidy.empty:
            write_docx_table(DESC_DOCX, desc_tidy, "Descriptives Table", subtitle_col="SDG")
        if not anova_tidy.empty:
            write_docx_table(ANOVA_DOCX, anova_tidy, "ANOVA Table", subtitle_col="SDG")
        if not tukey_tidy.empty:
            write_docx_table(TUKEY_DOCX, tukey_tidy, "Multiple Comparisons (Tukey HSD)", subtitle_col="SDG")
        print(f"save to docx：\n - {DESC_DOCX}\n - {ANOVA_DOCX}\n - {TUKEY_DOCX}")
    except Exception as e:
        print(f"no docx，use Excel：{e}")
        if not desc_tidy.empty:  desc_tidy.to_excel(DESC_XLSX, index=False)
        if not anova_tidy.empty: anova_tidy.to_excel(ANOVA_XLSX, index=False)
        if not tukey_tidy.empty: tukey_tidy.to_excel(TUKEY_XLSX, index=False)
        print(f"Excel output：\n - {DESC_XLSX}\n - {ANOVA_XLSX}\n - {TUKEY_XLSX}")

export_three_docs()
